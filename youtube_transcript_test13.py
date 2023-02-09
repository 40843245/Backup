from __future__ import absolute_import

from lxml import etree

import re

import docx

from docx.compat import Unicode
from docx.oxml import OxmlElement
from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.ns import NamespacePrefixedTag, nsmap, qn
from docx.shared import lazyproperty

from docx.oxml import ns



import json

import docx

from docx import shared

from docx import Document

from docx.shared import Inches

from docx import section

from docx.enum.section import WD_SECTION

from youtube_transcript_api import YouTubeTranscriptApi


from googletrans import Translator

import scrapetube





class YT_video():
    def ConcatentateListToLetter(self,otherOptions):
        insertedText=""
        if len(otherOptions)<0:
            insertedText=""
            raise "Too less options in the variable otherOptions!!!"
            
        if len(otherOptions)==0:
            insertedText="/"
            return insertedText
        if len(otherOptions)==1:
            insertedText=otherOptions[0]
            return insertedText
        
        for s in otherOptions:
            insertedText=insertedText+str(s)
        return insertedText
    
    def create_element(self,name):
        return OxmlElement(name)

    def create_attribute(self,element, name, value):
        element.set(ns.qn(name), value)


    def add_page_number(self,run,options):
        option0=options[0]
        # For other
        if option0!=0 and option0!=1:
            
            otherOptions=options[1:]
            insertedText=self.ConcatentateListToLetter(otherOptions)
            
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
            
            
            fldChar2 = self.create_element('w:t')
            fldChar2.text=insertedText
            self.create_attribute(fldChar2, 'w:fldCharType', 'separate')


            fldChar3 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')
            
            run._r.append(fldChar1)
            
            run._r.append(fldChar2)
            
            run._r.append(fldChar3)
            
        ### For current page and total page
        else:
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
            
            instrText = self.create_element('w:instrText')
            self.create_attribute(instrText, 'xml:space', 'preserve')
        
        
            if option0==0:
                instrText.text = "PAGE"
            elif option0==1:
                instrText.text = "NUMPAGES"
    
            fldChar2 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            
    def __init__(self,transcript_language):
        self.transcript_language=transcript_language
        self.YT_handler()
        
    def YT_handler(self):
        videos = scrapetube.get_channel("UCO-n4ZDDXKPKK29c5eaytpA")
        base_watch_video_url='https://www.youtube.com/watch?v='

        cnt=0
        for video in videos:
            cnt+=1
            v_url=video['videoId']
            video_url=str(base_watch_video_url)+v_url
            print(video_url)
            self.YT_Video_Transcript(video_url,'MiMi'+str(cnt),'.docx',['ja'])
            
    def YT_Video_Transcript(self,video_url,document_name,document_ext,transcript_language):
        translator = Translator()
        document = Document()
        
        self.document=document
        
        video_id=video_url.split('=')[1]
        print(video_id)
        video_info=YouTubeTranscriptApi.get_transcript(video_id,languages=transcript_language)
        print('-----------------------')
        
        line=int(input("How many line do you want to handle with?"))
        msg="The transcript of YT, link="+str(video_url)
        print(msg)
        
        
        p=self.document.add_paragraph("SHIT\n") 
        
        
        clr=docx.shared.RGBColor(255,255,0)
        
        print(dir(p.add_run(msg).font))
        print(type(p.add_run(msg).font))
        run = p.add_run(msg).font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        
        #run = p.add_run(msg).font.color = docx.enum.text.WD_COLOR_INDEX.YELLOW
        
        run = p.add_run(msg).font.color = clr
        
        
        
        #docx.dml.color.ColorFormat.type=docx.enum.dml.MSO_COLOR_TYPE.RGB
        
        
        print(clr)
    
        
        print(docx.dml.color.ColorFormat.type==docx.enum.dml.MSO_COLOR_TYPE.RGB)
        print(docx.dml.color.ColorFormat.type==docx.enum.dml.MSO_COLOR_TYPE.THEME)
        print(docx.dml.color.ColorFormat.type==docx.enum.dml.MSO_COLOR_TYPE.AUTO)
        
        #docx.enum.dml.MSO_COLOR_TYPE.RGB=clr
        
        
        
        print(docx.dml.color.ColorFormat.type==docx.enum.dml.MSO_COLOR_TYPE.RGB)
        print(docx.dml.color.ColorFormat.type==docx.enum.dml.MSO_COLOR_TYPE.THEME)
        print(docx.dml.color.ColorFormat.type==docx.enum.dml.MSO_COLOR_TYPE.AUTO)
        
        
        """
        print(type(docx.dml.color.ColorFormat.type))
        print(type(docx.dml.color.ColorFormat))
        print(type(docx.dml.color.ColorFormat.rgb))
        print(docx.dml.color.ColorFormat.rgb)
        print(docx.dml.color.ColorFormat.type)
        print(docx.dml.color.ColorFormat.type==docx.dml.color.ColorFormat.MSO_COLOR_TYPE.RGB)
        """
        #print(run.font.color.type)
        """
        document.add_page_break()
        count=0
        
        for v_info in video_info:
            count+=1
            for la in ['zh-tw','en','ja']:
                backup=v_info
                msg_para="The current language is "+str(la)
                print(msg_para)
                document.add_paragraph(msg_para)
                
                text=v_info['text']
                x=translator.translate(text, dest=la)
                
                backup['text']=x.text
                s=str(backup)
                print(s)
                document.add_paragraph(s)
            
            if count>=line:
                break
            
        self.document=document
        
        self.document=self.Auto_Add_Page_Number([1,0,""])
        
        self.document=self.Auto_Add_Page_Number([1,2,"/"])
        
        self.document=self.Auto_Add_Page_Number([1,1,""])
        
        """
        document=self.document
        
        self.document.save(document_name+str(document_ext))

    
    def Auto_Add_Page_Number(self,options):
        doc=self.document
        
        headerOrFooter=options[0]
        otherOptions=options[1:]
        if headerOrFooter==0:
            doc.sections[0].header.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            self.add_page_number(run=doc.sections[0].header.paragraphs[0].add_run(),options=otherOptions)
        elif headerOrFooter==1:
            doc.sections[0].footer.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            self.add_page_number(run=doc.sections[0].footer.paragraphs[0].add_run(),options=otherOptions)
        return doc

def main():
    YT_video(['ja'])
    

main()